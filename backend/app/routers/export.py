"""导出：Excel（.xlsx）与 Word（.docx），按分类层级组织，列序与主表一致。"""
import io

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db
from ..security import require_admin
from ..models import Classification, Indicator, IndicatorStatus, User

router = APIRouter(prefix="/export", tags=["导出"])

LEVELS = ["一级分类", "二级分类", "三级分类"]
HEADERS = ["来源标准/部分", *LEVELS, "标识符", "中文名称", "英文名称", "计量单位",
           "定义", "计算方法", "指标说明", "调查方法", "数据来源", "发布频率"]


def _walk(db: Session, parent_id=None, path=None):
    """深度优先遍历分类树，产出 (node, [一级,二级,三级名称]) 序列。"""
    path = path or []
    nodes = (db.query(Classification).filter(Classification.parent_id == parent_id)
             .order_by(Classification.sort_order, Classification.id).all())
    for n in nodes:
        p = path + [n.name]
        yield n, p
        yield from _walk(db, n.id, p)


def _indicators(db: Session, class_id: int):
    return (db.query(Indicator)
            .filter(Indicator.classification_id == class_id, Indicator.status == IndicatorStatus.active)
            .order_by(Indicator.identifier).all())


@router.get("/excel", summary="导出 Excel")
def export_excel(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook(); ws = wb.active; ws.title = "卫生统计指标"
    thin = Side(style="thin", color="D0D0D0"); border = Border(thin, thin, thin, thin)
    ws.append(HEADERS)
    for c in range(1, len(HEADERS) + 1):
        cell = ws.cell(1, c)
        cell.font = Font(name="宋体", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E5F")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    for node, path in _walk(db):
        for ind in _indicators(db, node.id):
            levels = (path + ["", "", ""])[:3]
            ws.append([
                ind.source_standard.title if ind.source_standard else "",
                levels[0], levels[1], levels[2],
                ind.identifier, ind.name_cn, ind.name_en, ind.unit,
                ind.definition, ind.method, ind.description,
                ind.survey_method, ind.data_source, ind.frequency,
            ])
            r = ws.max_row
            for c in range(1, len(HEADERS) + 1):
                cc = ws.cell(r, c)
                cc.font = Font(name="宋体", size=10)
                cc.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                cc.border = border

    widths = [28, 12, 12, 12, 14, 22, 30, 8, 40, 32, 40, 10, 16, 10]
    from openpyxl.utils import get_column_letter
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=health_indicators.xlsx"})


@router.get("/word", summary="导出 Word")
def export_word(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # 每个指标按此顺序，逐行输出（左列字段名，右列内容）
    FIELDS = [
        ("来源标准/部分", lambda i: i.source_standard.title if i.source_standard else ""),
        ("标识符", lambda i: i.identifier),
        ("中文名称", lambda i: i.name_cn),
        ("英文名称", lambda i: i.name_en),
        ("计量单位", lambda i: i.unit),
        ("定义", lambda i: i.definition),
        ("计算方法", lambda i: i.method),
        ("指标说明", lambda i: i.description),
        ("调查方法", lambda i: i.survey_method),
        ("数据来源", lambda i: i.data_source),
        ("发布频率", lambda i: i.frequency),
    ]

    def style_run(run, size=10.5, bold=False, color=None):
        """中文宋体、英文 Times New Roman。"""
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    def write_cell(cell, text, bold=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run("" if text is None else str(text))
        style_run(run, bold=bold)

    doc = Document()
    # 文档默认样式：中文宋体、英文 Times New Roman
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    title = doc.add_heading(level=0)
    style_run(title.add_run("卫生统计指标（含元数据）"), size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x5F))

    def add_heading(text, level):
        h = doc.add_heading(level=level)
        style_run(h.add_run(text), size={1: 15, 2: 13, 3: 12, 4: 12}.get(level, 12),
                  bold=True, color=RGBColor(0x0F, 0x76, 0x6E))

    def emit(parent_id=None, depth=1):
        nodes = (db.query(Classification).filter(Classification.parent_id == parent_id)
                 .order_by(Classification.sort_order, Classification.id).all())
        for n in nodes:
            add_heading(n.name, min(depth, 4))
            for ind in _indicators(db, n.id):
                tbl = doc.add_table(rows=0, cols=2)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                for label, getter in FIELDS:
                    cells = tbl.add_row().cells
                    write_cell(cells[0], label, bold=True)
                    write_cell(cells[1], getter(ind))
                    cells[0].width = Cm(3.2)
                    cells[1].width = Cm(13.8)
                doc.add_paragraph()  # 指标间留白
            emit(n.id, depth + 1)

    emit()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=health_indicators.docx"})
